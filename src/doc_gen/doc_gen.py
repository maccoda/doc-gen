"""
Doc-Gen
Documentation generator from markdown files.
"""
import sys
import getopt
import re
import os

import mistune
from doc_elements import Heading, Paragraph, BoldText, ItalicText, Text, DocumentElement
from docx import Document


class IdentityRenderer(mistune.Renderer):
    """
    Renderer which just returns a typed object of the arguments for later
    processing to construct the DOCX document
    """

    def __init__(self, **kwargs):
        super().__init__()

    def placeholder(self):
        return []

    def header(self, text, level, raw=None):
        return [Heading(level, text)]

    def paragraph(self, text):
        return [Paragraph(text)]

    def double_emphasis(self, text):
        return [BoldText(text)]

    def emphasis(self, text):
        return [ItalicText(text)]

    def text(self, text):
        return [Text(text)]


class DocumentBuilder:
    """Builder of the DOCX document from DocumentElements"""

    def __init__(self, md_dir, templ_name=None):
        self.document = Document()
        self.elements = {}
        doc_render = IdentityRenderer()
        # Read all md files in the directory
        for cur_file in os.listdir(md_dir):
            if cur_file.endswith('.md'):
                with open(md_dir + '/' + cur_file) as in_file:
                    # Parse and convert the md files
                    elem_list = mistune.Markdown(
                        renderer=doc_render).output(in_file.read())
                    head = elem_list[0]
                    # Find the key for the content
                    if isinstance(head, Heading):
                        key = head.heading_text()
                        # Store the remaining elements
                        self.elements[key.as_str()] = elem_list[1:]
                    else:
                        print("Top element of the Markdown files must be a heading")
                        raise Exception
        if templ_name:
            self.template_doc = Document(templ_name)
        else:
            self.template_doc = None

    def create(self, out_file_name):
        """
        Creates and saves a DOCX file from its current elements adding them to
        when their appropriate heading is reached. If no template was provided
        it will simply generate the content onto a blank document.

        TODO How will it function when no template is there and there are
        multiple element lists???
        """
        if self.template_doc:
            self._populate_with_template()
        # HACK for non-template scenario
        # Forcing all of the files to start with Content Header
        else:
            for item in self.elements['Content']:
                if isinstance(item, DocumentElement):
                    # Need to have as -1 to allow for first heading to be
                    # heading 1 and match the design convention
                    item.append_to_document(self.document, -1)
                else:
                    raise Exception

        self.document.save(out_file_name)

    def _populate_with_template(self):
        """Populates the document with the template content"""
        heading_match = r"Heading (\d+)"
        for elem in self.template_doc.paragraphs:
            style_name = elem.style.name
            style = self.document.styles[style_name]
            para = self.document.add_paragraph(elem.text)
            para.style = style
            # Check if should add under this
            if elem.text in self.elements:
                parent_level = re.match(heading_match, style_name).group(1)
                for item in self.elements[elem.text]:
                    if isinstance(item, DocumentElement):
                        item.append_to_document(
                            self.document, int(parent_level) - 1)
                    else:
                        raise Exception


def print_help():
    """Print help message"""
    print("doc_gen.py -i <input> -o <output>")
    print("-i,--input\t name of directory containing Markdown files to convert")
    print("-o,--output\t name to write output DOCX file to (with file extension)")


def parse_args(args):
    """Parse the provided arguments and return a dictionary of received arguments"""
    try:
        opts, args = getopt.getopt(
            args, "hi:o:t:", ["input", "output", "help", "template"])
    except getopt.GetoptError:
        print_help()
        sys.exit(2)

    returned = {}
    for opt, arg in opts:
        if opt in ('-h', '--help'):
            print_help()
            sys.exit()
        elif opt in ('-i', '--input'):
            returned['input_dir'] = arg
        elif opt in ('-o', '--output'):
            returned['output_file'] = arg
        elif opt in ('-t', '--template'):
            returned['template_file'] = arg

    return returned


def get_args(args):
    """Parse and check correct arguments present"""
    parsed = parse_args(args)
    if 'input_dir' not in parsed or 'output_file' not in parsed:
        print("Error: Not all arguments provided")
        print_help()
        sys.exit(2)

    return (parsed['input_dir'], parsed['output_file'], parsed.get('template_file'))


def main(input_dir, output_file_name, template_name):
    """Main method taking the markdown file and creating DOCX file"""
    doc_builder = DocumentBuilder(input_dir, template_name)
    doc_builder.create(output_file_name)



if __name__ == '__main__':
    m_input_dir, m_output_file_name, m_template_name = get_args(sys.argv[1:])
    main(m_input_dir, m_output_file_name, m_template_name)
