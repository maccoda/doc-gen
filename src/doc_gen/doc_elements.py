"""
Module containing all of parsed elements from the Markdown file and the relevant
mappings to the DOCX types
"""
import os
import docx


class DocumentElement:
    """Single DOCX document element. Capable of appending itself to a document"""

    def append_to_document(self, document, parent_level):
        """
        Append this element to the document with all heading levels being
        under the one provided.
        """
        pass


class SingleElementTextWrapper(DocumentElement):
    """Element expecting only a single Text element to be within its content"""

    def __init__(self, text):
        # Expect only a single Text element
        if len(text) > 1:
            raise Exception(
                "Single element text wrapper with more than single element")
        self.text = text[0]


class StyledText(SingleElementTextWrapper):
    """Text element with specialized styling. e.g. Bold & Italic"""

    def append_to_document(self, document, parent_level):
        run = document.paragraphs[len(document.paragraphs) -
                                  1].add_run(self.text.as_str())
        self.add_style_to_run(run)

    def add_style_to_run(self, run):
        """Add styling to the provided Run element"""
        pass


class Heading(SingleElementTextWrapper):
    """Heading element in document. Contains text and the level of heading"""

    def __init__(self, level, text):
        self.level = level
        super().__init__(text)

    def heading_text(self):
        """Return the heading text"""
        return self.text

    def append_to_document(self, document, parent_level):
        document.add_heading(self.text.as_str(),
                             level=(self.level + parent_level))


class Paragraph(DocumentElement):
    """
    Single paragraph element. This is a glob of text that is expected to be
    separated by a new line, mapping directly to a Paragraph element in the
    Document
    """

    def __init__(self, text_elements):
        self.text_elements = text_elements

    def append_to_document(self, document, parent_level):
        # Add a new paragraph
        document.add_paragraph('')
        for item in self.text_elements:
            item.append_to_document(document, parent_level)


class BoldText(StyledText):
    """Bold text element"""

    def add_style_to_run(self, run):
        run.bold = True


class ItalicText(StyledText):
    """Italic text element"""

    def add_style_to_run(self, run):
        run.italic = True


class Text(DocumentElement):
    """Basic text element, containing only plain string"""

    def __init__(self, text):
        # Trim new lines as they have no value in Markdown but affect docx
        self.text = str(text).replace('\n', ' ')

    def append_to_document(self, document, parent_level):
        document.paragraphs[len(document.paragraphs) -
                            1].add_run(self.text)

    def as_str(self):
        """Returns the string representation of the this Text element"""
        return self.text


class Image(DocumentElement):
    """Image element"""

    def __init__(self, image_path):
        self.path = os.path.abspath(image_path)

    def append_to_document(self, document, parent_level):
        document.add_picture(self.path)


class Link(SingleElementTextWrapper):
    """Hyperlink element"""

    def __init__(self, url, text):
        self.url = url
        super().__init__(text)

    def append_to_document(self, document, parent_level):
        paragraph = document.paragraphs[len(document.paragraphs) - 1]

        part = paragraph.part
        r_id = part.relate_to(
            self.url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id,)

        new_run = docx.oxml.shared.OxmlElement('w:r')
        rPr = docx.oxml.shared.OxmlElement('w:rPr')

        new_run.append(rPr)
        new_run.text = self.text.as_str()
        hyperlink.append(new_run)

        paragraph._p.append(hyperlink)


class WrittenList(DocumentElement):
    """Ordered or unordered list containing collection of ListElements"""

    def __init__(self, elements, ordered):
        self.elements = elements
        self.ordered = ordered

    def append_to_document(self, document, parent_level):
        for elem in self.elements:
            paragraph = document.add_paragraph('')
            if self.ordered:
                paragraph.style = 'List Number'
            else:
                paragraph.style = 'List Bullet'
            elem.append_to_document(document, parent_level)


class ListElement(DocumentElement):
    """Single element in a list"""

    def __init__(self, text):
        # Element does not need to be text, may have style or link
        self.text = text

    def append_to_document(self, document, parent_level):
        if len(self.text) > 1:
            print("List element with more than one item")
            print(self.text)
        self.text[0].append_to_document(document, parent_level)


class Table(DocumentElement):
    """Root element for a table"""

    def __init__(self, header, content):
        self.header = header[0]
        self.content = content

    def append_to_document(self, document, parent_level):
        if isinstance(self.header, TableRow):
            cols = self.header.number_of_columns()
            table = document.add_table(0, cols, 'Table Grid')
            # Add the header
            header_row = self.header.add_heading_row(table)
            for elem in self.content:
                if isinstance(elem, TableRow):
                    elem.add_row(table)
                else:
                    raise Exception('This table is not filled with rows...')
        else:
            raise Exception('This table has headers that are not a table row')


class TableRow(DocumentElement):
    """Single row of a table"""

    def __init__(self, elements):
        self.elements = elements

    def number_of_columns(self):
        """Returns the number of columns in this row"""
        return len(self.elements)

    def add_heading_row(self, table):
        """Adds a heading row to the provided table"""
        self._write_row(table, True)

    def add_row(self, table):
        """Adds a text row to the provided table"""
        self._write_row(table, False)

    def _write_row(self, table, heading):
        """
        Writes a row to the provided table and uses heading boolean to determine if should be
        written with heading style or not.
        """
        if heading:
            style = 'Strong'
        else:
            style = None
        row = table.add_row()
        for i, cell in enumerate(self.elements):
            if isinstance(cell, TableCell):
                cell.write_to_cell(row.cells[i], style)
        return row


class TableCell(SingleElementTextWrapper):
    """Singular cell within a table"""

    def __init__(self, text):
        super().__init__(text)

    def write_to_cell(self, cell, style):
        """Write self to the cell provided with the given style"""
        cell.text = self.text.as_str()
        if style:
            cell.paragraphs[0].runs[0].style = style
