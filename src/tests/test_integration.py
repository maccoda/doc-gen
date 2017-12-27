"""Integration tests"""

import sys
import os
from docx import Document


def rel_to_abs_path(relative):
    """Returns the absolute path of the file given as a relative path"""
    return os.path.abspath(os.path.join(os.path.dirname(__file__), relative))


PATH = rel_to_abs_path('../doc_gen')
sys.path.insert(0, PATH)
import doc_gen

DEF_STYLE = 'Default Paragraph Font'
ITALIC_RUN = 'italic'
BOLD_RUN = 'bold'


def test_conversion_from_markdown():
    """High level integration test"""
    in_path = rel_to_abs_path('./single')
    out_path = rel_to_abs_path('./output.docx')
    doc_gen.main(in_path, out_path, None)
    # Now to test
    document = Document(out_path)
    paragraphs = document.paragraphs
    assert len(paragraphs) == NUM_PARAS_TEST1

    assertions_of_test1md(paragraphs, 0, 1)

    # Clean Up
    os.remove(out_path)


def test_addition_to_template():
    """High level test using template"""
    template_name = generate_template_document()
    in_path = rel_to_abs_path('./single')
    out_path = rel_to_abs_path('./output.docx')
    template_path = rel_to_abs_path('../../' + template_name)
    doc_gen.main(in_path, out_path, template_path)

    paragraphs = Document(out_path).paragraphs
    assert len(paragraphs) == 7 + NUM_PARAS_TEST1
    # Template document
    assert_heading(paragraphs[0], 'Template Heading', 'Heading 1')
    assert_heading(paragraphs[1], 'References', 'Heading 2')
    assert_text_runs(paragraphs[2], 'Reference 1 - Something')
    assert_text_runs(paragraphs[3], 'Reference 2 - Another thing')
    assert_heading(paragraphs[4], 'Content', 'Heading 2')

    # Added content asserts
    assertions_of_test1md(paragraphs, 5, 3)

    # Template after content
    assert_heading(paragraphs[5 + NUM_PARAS_TEST1 + 0],
                   'Post Content Information', 'Heading 2')
    assert_heading(paragraphs[5 + NUM_PARAS_TEST1 + 1], 'Finishing Touches', 'Heading 3')

    # Clean up
    os.remove(template_path)
    os.remove(out_path)


def test_mutli_addition_to_template():
    """
    Testing the final form of generator, whereby there is a template generated
    that has several sections to be filled and it is filled by the markdown
    files within a certain directory.
    """
    template_name = generate_multi_section_template()
    in_path = rel_to_abs_path('./multi')
    out_path = rel_to_abs_path('./output.docx')
    template_path = rel_to_abs_path('../../' + template_name)
    doc_gen.main(in_path, out_path, template_path)

    paragraphs = Document(out_path).paragraphs
    assert len(paragraphs) == 10 + NUM_PARAS_TEST1
    # Template document
    paras = ParagraphIterator(paragraphs)
    assert_heading(paras.next(), 'Template Heading', 'Heading 1')
    assert_heading(paras.next(), 'Overview', 'Heading 2')
    # Added Overview asserts
    assert_text_runs(
        paras.next(), 'A brief overview of the document that we have.',
        [DEF_STYLE, ITALIC_RUN, DEF_STYLE])

    assert_heading(paras.next(), 'References', 'Heading 2')
    assert_text_runs(paras.next(), 'Reference 1 - Something')
    assert_text_runs(paras.next(), 'Reference 2 - Another thing')
    assert_heading(paras.next(), 'Content', 'Heading 2')

    # Added content asserts
    assert_text_runs(paras.next(), 'Some content about the topic needed to be discussed.',
        [DEF_STYLE, BOLD_RUN, DEF_STYLE])
    assert_heading(paras.next(), 'Subsection', 'Heading 3')
    assert_text_runs(paras.next(), 'This is a subsection of the content')
    assert_text_runs(paras.next(), 'With an image of the documentation rendered')
    # Empty line
    assert_empty_line(paras.next())
    # Image
    assert_image(paras.next())



    # Template after content
    assert_heading(paras.next(),
                   'Post Content Information', 'Heading 2')
    assert_text_runs(paras.next(), 'This is some information after the fact')
    assert_heading(paras.next(),
                   'Finishing Touches', 'Heading 3')

    # Clean up
    os.remove(template_path)
    os.remove(out_path)

# Number of paragraphs to be generated from test1.md
NUM_PARAS_TEST1 = 6

def assertions_of_test1md(paragraphs, start_index, start_heading_level):
    """
    Assertions for the content generated from test1.md. Takes the list of
    paragraphs to check and the starting index in the list for the generated
    content. The heading level defines the expected starting level which should
    be one below the parent section.
    """
    # Check the headings
    assert_heading(paragraphs[start_index + 0], 'Heading 1', 'Heading ' + str(start_heading_level))
    assert_heading(paragraphs[start_index + 1], 'Heading 2',
                   'Heading ' + str(start_heading_level + 1))
    assert_heading(paragraphs[start_index + 4],
                   'Another Heading 2', 'Heading ' + str(start_heading_level + 1))

    # Check the basic paragraphs/text
    assert_text_runs(paragraphs[start_index + 3], 'Now a second paragraph')
    assert_text_runs(paragraphs[start_index + 5], 'Let us add some more text.')
    assert_text_runs(paragraphs[start_index + 2], 'Some paragraph text with bold and italics.',
                     [DEF_STYLE, BOLD_RUN, DEF_STYLE, ITALIC_RUN, DEF_STYLE])


def generate_template_document():
    """Generates a mock template DOCX document. Returning the file name"""
    document = Document()
    document.add_heading('Template Heading', 1)
    document.add_heading('References', 2)
    document.add_paragraph('Reference 1 - Something')
    document.add_paragraph('Reference 2 - Another thing')
    document.add_heading('Content', 2)
    document.add_heading('Post Content Information', 2)
    document.add_heading('Finishing Touches', 3)

    doc_name = './Template.docx'
    document.save(doc_name)
    return doc_name


class ParagraphIterator:
    """Iterator over the document paragraphs to remove hardcoded indexes"""

    def __init__(self, paras):
        self.curr = 0
        self.paragraphs = paras

    def next(self):
        """Get next paragraph"""
        result = self.paragraphs[self.curr]
        self.curr += 1
        return result


def generate_multi_section_template():
    """
    Generates a mock template DOCX document with sections to be filled being the
    Content and Overview sections.

    Returning the file name
    """
    document = Document()
    document.add_heading('Template Heading', 1)
    document.add_heading('Overview', 2)
    document.add_heading('References', 2)
    document.add_paragraph('Reference 1 - Something')
    document.add_paragraph('Reference 2 - Another thing')
    document.add_heading('Content', 2)
    document.add_heading('Post Content Information', 2)
    document.add_paragraph('This is some information after the fact')
    document.add_heading('Finishing Touches', 3)

    doc_name = './Template_multi.docx'
    document.save(doc_name)
    return doc_name


def assert_heading(heading, text, style_name):
    """Assert that the provided heading matches the text and style name provided"""
    assert heading.text == text
    assert heading.style.name == style_name


def assert_text_runs(text_run, text, run_styles=None):
    """Assert that the provided text has the content and styles provided"""
    assert text_run.text == text
    if not run_styles:
        assert len(text_run.runs) == 1
        assert text_run.runs[0].style.name == DEF_STYLE
    else:
        assert len(text_run.runs) == len(run_styles)

def assert_image(image):
    """Assert that the provided element is an image"""
    # Don't really know what is going on here
    assert len(image.runs) == 1

def assert_empty_line(element):
    """Assert that the provided element is an empty line"""
    assert element.text == ''
    assert len(element.runs) == 0
